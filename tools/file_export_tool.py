import os
from loguru import logger


def export_as_txt(content: str, file_name: str = "result") -> str:
    os.makedirs("vector_db", exist_ok=True)
    output_path = f"vector_db/{file_name}_export.txt"

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)

    logger.info(f"Exported as TXT: {output_path}")
    return output_path


def export_as_docx(content: str, file_name: str = "result") -> str:
    from docx import Document

    os.makedirs("vector_db", exist_ok=True)
    output_path = f"vector_db/{file_name}_export.docx"

    doc = Document()
    doc.add_heading("DocuMind AI — Generated Document", 0)

    for line in content.split("\n"):
        if line.strip():
            if line.startswith("##"):
                doc.add_heading(line.replace("##", "").strip(), 2)
            elif line.startswith("#"):
                doc.add_heading(line.replace("#", "").strip(), 1)
            else:
                doc.add_paragraph(line)

    doc.save(output_path)
    logger.info(f"Exported as DOCX: {output_path}")
    return output_path


def export_as_csv(content: str, file_name: str = "result") -> str:
    import pandas as pd

    os.makedirs("vector_db", exist_ok=True)
    output_path = f"vector_db/{file_name}_export.csv"

    lines = [
        line.strip()
        for line in content.split("\n")
        if line.strip()
    ]
    df = pd.DataFrame(lines, columns=["Content"])
    df.to_csv(output_path, index=False)

    logger.info(f"Exported as CSV: {output_path}")
    return output_path